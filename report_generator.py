from typing import List, Dict
from datetime import datetime
import os
from dotenv import load_dotenv

load_dotenv()

try:
    import google.genai as genai
    GENAI_AVAILABLE = True
except ImportError:
    genai = None
    GENAI_AVAILABLE = False

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if GEMINI_API_KEY and GENAI_AVAILABLE:
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        USE_GEMINI = True
    except Exception as e:
        print(f"Gemini not available: {e}")
        USE_GEMINI = False
        client = None
else:
    USE_GEMINI = False
    client = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None

class ResearchReportGenerator:
    def __init__(self):
        self.client = client
        print("[OK] Using Gemini for report generation")

    def generate_report(self, query: str, synthesis_result: Dict,
                       agreements: List[Dict], contradictions: List[Dict],
                       documents: List[str]) -> str:
        """Generate comprehensive research report"""
        
        prompt = f"""Create a comprehensive research report based on the following analysis.

Research Question: {query}

Synthesis:
{synthesis_result.get("synthesis", "")}

Agreements Between Documents:
{chr(10).join([str(a) for a in agreements]) if agreements else "No agreements found"}

Contradictions Found:
{chr(10).join([str(c) for c in contradictions]) if contradictions else "No contradictions found"}

Structure the report as:
1. EXECUTIVE SUMMARY (2-3 sentences)
2. RESEARCH QUESTION
3. KEY FINDINGS (bullet points)
4. CROSS-DOCUMENT ANALYSIS
5. AREAS OF CONSENSUS
6. AREAS OF DISAGREEMENT
7. RESEARCH GAPS
8. CONCLUSION
9. REFERENCES

Make it professional and suitable for publication."""
        
        if USE_GEMINI and self.client:
            try:
                response = self.client.models.generate_content(
                    model='gemini-2.0-flash-thinking-exp-1219',
                    contents=prompt
                )
                return response.text if response and response.text else "Error generating report"
            except Exception as e:
                return f"Error generating report: {str(e)}"
        else:
            return "No LLM provider configured"

    def export_to_docx(self, report_content: str, query: str,
                      documents: List[str], output_path: str):
        """Export report to Word document"""
        
        if Document is None:
            print("python-docx not available")
            return
        
        try:
            doc = Document()
            
            title = doc.add_heading("Research Report", 0)
            if WD_ALIGN_PARAGRAPH:
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading("Research Question", level=1)
            doc.add_paragraph(query)
            
            doc.add_heading("Documents Analyzed", level=1)
            for doc_name in documents:
                doc.add_paragraph(doc_name, style='List Bullet')
            
            doc.add_heading("Report", level=1)
            doc.add_paragraph(report_content)
            
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if WD_ALIGN_PARAGRAPH:
                footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(output_path)
            print(f"Report saved to: {output_path}")
        except Exception as e:
            print(f"Error exporting to docx: {e}")

    def export_to_markdown(self, report_content: str, query: str,
                          documents: List[str], output_path: str):
        """Export report to Markdown"""
        
        markdown_content = f"""# Research Report

## Research Question
{query}

## Documents Analyzed
{chr(10).join([f"- {doc}" for doc in documents])}

## Report
{report_content}

---
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        try:
            with open(output_path, 'w') as f:
                f.write(markdown_content)
            print(f"Report saved to: {output_path}")
        except Exception as e:
            print(f"Error exporting to markdown: {e}")

    def generate_analytics(self, synthesis_results: List[Dict],
                          documents: List[str]) -> Dict:
        """Generate analytics on document relevance and coverage"""
        
        doc_citation_count = {}
        for doc in documents:
            doc_citation_count[doc] = 0
        
        for result in synthesis_results:
            for doc in result.get("documents_cited", []):
                if doc in doc_citation_count:
                    doc_citation_count[doc] += 1
        
        total_citations = sum(doc_citation_count.values())
        
        analytics = {
            "total_documents": len(documents),
            "total_queries": len(synthesis_results),
            "document_relevance": {
                doc: {
                    "citation_count": count,
                    "relevance_percentage": (count / total_citations * 100) if total_citations > 0 else 0
                }
                for doc, count in doc_citation_count.items()
            },
            "most_relevant_documents": sorted(
                doc_citation_count.items(),
                key=lambda x: x[1],
                reverse=True
            )[:5]
        }
        
        return analytics
